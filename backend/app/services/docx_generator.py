from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.services.resume_templates import get_template


class DOCXResumeGenerator:
    """
    ATS-friendly resume DOCX generator.

    Supported templates:
        - single_column
        - double_column

    The double-column template follows the structure
    of the supplied reference resume:
        Header
        Career Objective
        Left sidebar
        Main career content
    """

    # ==========================================================
    # PUBLIC ENTRY POINT
    # ==========================================================

    @staticmethod
    def generate(
        generated_resume: Dict,
        output_path: str,
        template_key: str = "single_column",
        mailing_address: Optional[str] = None,
    ) -> str:

        template = get_template(template_key)

        output = Path(output_path)

        output.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()

        DOCXResumeGenerator._configure_page(
            document
        )

        DOCXResumeGenerator._configure_styles(
            document
        )

        if template.key == "single_column":

            DOCXResumeGenerator._build_single_column(
                document,
                generated_resume,
            )

        elif template.key == "double_column":

            if not mailing_address:
                raise ValueError(
                    "Mailing address is required "
                    "for the double-column template."
                )

            DOCXResumeGenerator._build_double_column(
                document,
                generated_resume,
                mailing_address,
            )

        else:

            raise ValueError(
                f"Unsupported template: {template_key}"
            )

        document.save(
            str(output)
        )

        return str(output)

    # ==========================================================
    # DOCUMENT CONFIGURATION
    # ==========================================================

    @staticmethod
    def _configure_page(
        document: Document,
    ) -> None:

        section = document.sections[0]

        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.2)

    @staticmethod
    def _configure_styles(
        document: Document,
    ) -> None:

        normal = document.styles["Normal"]

        normal.font.name = "Arial"
        normal.font.size = Pt(9)

        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.line_spacing = 1.0

        # Ensure Arial is applied to East Asian fonts too.
        normal._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial",
        )

    # ==========================================================
    # SINGLE COLUMN TEMPLATE
    # ==========================================================

    @staticmethod
    def _build_single_column(
        document: Document,
        resume: Dict,
    ) -> None:

        DOCXResumeGenerator._add_name_header(
            document,
            resume,
        )

        DOCXResumeGenerator._add_contact_information(
            document,
            resume,
        )

        DOCXResumeGenerator._add_section(
            document,
            "Professional Summary",
        )

        summary = resume.get(
            "summary",
            "",
        )

        if summary:

            DOCXResumeGenerator._add_body_text(
                document,
                summary,
            )

        DOCXResumeGenerator._add_skills(
            document,
            resume,
        )

        DOCXResumeGenerator._add_experience(
            document,
            resume,
        )

        DOCXResumeGenerator._add_projects(
            document,
            resume,
        )

        DOCXResumeGenerator._add_education(
            document,
            resume,
        )

        DOCXResumeGenerator._add_certifications(
            document,
            resume,
        )

        DOCXResumeGenerator._add_links(
            document,
            resume,
        )

    # ==========================================================
    # DOUBLE COLUMN TEMPLATE
    # ==========================================================

    @staticmethod
    def _build_double_column(
        document: Document,
        resume: Dict,
        mailing_address: str,
    ) -> None:

        # ------------------------------------------------------
        # TOP HEADER
        # ------------------------------------------------------

        DOCXResumeGenerator._add_name_header(
            document,
            resume,
        )

        DOCXResumeGenerator._add_contact_information(
            document,
            resume,
        )

        # ------------------------------------------------------
        # CAREER OBJECTIVE
        # ------------------------------------------------------

        DOCXResumeGenerator._add_section(
            document,
            "Career Objective",
        )

        objective = DOCXResumeGenerator._create_objective(
            resume
        )

        DOCXResumeGenerator._add_body_text(
            document,
            objective,
        )

        # ------------------------------------------------------
        # START TWO-COLUMN AREA
        # ------------------------------------------------------

        DOCXResumeGenerator._enable_two_columns(
            document
        )

        # ======================================================
        # LEFT COLUMN
        # ======================================================

        DOCXResumeGenerator._add_section(
            document,
            "Mailing Address",
        )

        DOCXResumeGenerator._add_multiline_text(
            document,
            mailing_address,
        )

        DOCXResumeGenerator._add_skills(
            document,
            resume,
        )

        DOCXResumeGenerator._add_education(
            document,
            resume,
        )

        DOCXResumeGenerator._add_certifications(
            document,
            resume,
        )

        DOCXResumeGenerator._add_links(
            document,
            resume,
        )

        # ======================================================
        # COLUMN BREAK
        # ======================================================

        DOCXResumeGenerator._add_column_break(
            document
        )

        # ======================================================
        # RIGHT COLUMN
        # ======================================================

        DOCXResumeGenerator._add_experience(
            document,
            resume,
        )

        DOCXResumeGenerator._add_projects(
            document,
            resume,
        )

    # ==========================================================
    # HEADER
    # ==========================================================

    @staticmethod
    def _add_name_header(
        document: Document,
        resume: Dict,
    ) -> None:

        name = resume.get(
            "name",
            "",
        )

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            name
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(18)

    @staticmethod
    def _add_contact_information(
        document: Document,
        resume: Dict,
    ) -> None:

        values = []

        email = resume.get(
            "email",
            "",
        )

        phone = resume.get(
            "phone",
            "",
        )

        location = resume.get(
            "location",
            "",
        )

        if email:
            values.append(email)

        if phone:
            values.append(phone)

        if location:
            values.append(location)

        if values:

            paragraph = document.add_paragraph()

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(
                " | ".join(values)
            )

            run.font.name = "Arial"
            run.font.size = Pt(8.5)

    @staticmethod
    def _add_links(
        document: Document,
        resume: Dict,
    ) -> None:

        links = resume.get(
            "links",
            [],
        )

        if not links:
            return

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            " | ".join(links)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)

    # ==========================================================
    # SECTION HEADING
    # ==========================================================

    @staticmethod
    def _add_section(
        document: Document,
        title: str,
    ) -> None:

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            title.upper()
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

        # Add bottom border.
        p = paragraph._p
        pPr = p.get_or_add_pPr()

        pBdr = OxmlElement(
            "w:pBdr"
        )

        bottom = OxmlElement(
            "w:bottom"
        )

        bottom.set(
            qn("w:val"),
            "single",
        )

        bottom.set(
            qn("w:sz"),
            "6",
        )

        bottom.set(
            qn("w:space"),
            "1",
        )

        bottom.set(
            qn("w:color"),
            "000000",
        )

        pBdr.append(
            bottom
        )

        pPr.append(
            pBdr
        )

    # ==========================================================
    # BODY TEXT
    # ==========================================================

    @staticmethod
    def _add_body_text(
        document: Document,
        text: str,
    ) -> None:

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run(
            text
        )

        run.font.name = "Arial"
        run.font.size = Pt(9)

    @staticmethod
    def _add_multiline_text(
        document: Document,
        text: str,
    ) -> None:

        for line in text.splitlines():

            line = line.strip()

            if not line:
                continue

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(
                line
            )

            run.font.name = "Arial"
            run.font.size = Pt(8.5)

    # ==========================================================
    # SKILLS
    # ==========================================================

    @staticmethod
    def _add_skills(
        document: Document,
        resume: Dict,
    ) -> None:

        skills = resume.get(
            "skills",
            [],
        )

        if not skills:
            return

        DOCXResumeGenerator._add_section(
            document,
            "Technical Skills",
        )

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(3)

        run = paragraph.add_run(
            ", ".join(skills)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8.5)

    # ==========================================================
    # EXPERIENCE
    # ==========================================================

    @staticmethod
    def _add_experience(
        document: Document,
        resume: Dict,
    ) -> None:

        experience = resume.get(
            "experience",
            [],
        )

        if not experience:
            return

        DOCXResumeGenerator._add_section(
            document,
            "Experience",
        )

        for item in experience:

            role = item.get(
                "role",
                "",
            )

            company = item.get(
                "company",
                "",
            )

            duration = item.get(
                "duration",
                "",
            )

            heading_values = []

            if role:
                heading_values.append(role)

            if company:
                heading_values.append(company)

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(
                " - ".join(
                    heading_values
                )
            )

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)

            if duration:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    duration
                )

                run.italic = True
                run.font.size = Pt(8)

            for bullet in item.get(
                "bullets",
                [],
            ):

                DOCXResumeGenerator._add_bullet(
                    document,
                    bullet,
                )

    # ==========================================================
    # PROJECTS
    # ==========================================================

    @staticmethod
    def _add_projects(
        document: Document,
        resume: Dict,
    ) -> None:

        projects = resume.get(
            "projects",
            [],
        )

        if not projects:
            return

        DOCXResumeGenerator._add_section(
            document,
            "Projects",
        )

        for project in projects:

            name = project.get(
                "name",
                "",
            )

            if name:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    name
                )

                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9)

            description = project.get(
                "description",
                "",
            )

            if description:

                DOCXResumeGenerator._add_body_text(
                    document,
                    description,
                )

            technologies = project.get(
                "technologies",
                [],
            )

            if technologies:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    "Technologies: "
                    + ", ".join(
                        technologies
                    )
                )

                run.italic = True
                run.font.size = Pt(8)

            for bullet in project.get(
                "bullets",
                [],
            ):

                DOCXResumeGenerator._add_bullet(
                    document,
                    bullet,
                )

    # ==========================================================
    # EDUCATION
    # ==========================================================

    @staticmethod
    def _add_education(
        document: Document,
        resume: Dict,
    ) -> None:

        education = resume.get(
            "education",
            [],
        )

        if not education:
            return

        DOCXResumeGenerator._add_section(
            document,
            "Education",
        )

        for item in education:

            degree = item.get(
                "degree",
                "",
            )

            institution = item.get(
                "institution",
                "",
            )

            duration = item.get(
                "duration",
                "",
            )

            if degree:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(
                    degree
                )

                run.bold = True
                run.font.size = Pt(8.5)

            if institution:

                paragraph = document.add_paragraph(
                    institution
                )

                paragraph.paragraph_format.space_after = Pt(1)

            if duration:

                paragraph = document.add_paragraph(
                    duration
                )

                paragraph.paragraph_format.space_after = Pt(2)

    # ==========================================================
    # CERTIFICATIONS
    # ==========================================================

    @staticmethod
    def _add_certifications(
        document: Document,
        resume: Dict,
    ) -> None:

        certifications = resume.get(
            "certifications",
            [],
        )

        if not certifications:
            return

        DOCXResumeGenerator._add_section(
            document,
            "Certifications",
        )

        for certification in certifications:

            DOCXResumeGenerator._add_bullet(
                document,
                certification,
            )

    # ==========================================================
    # BULLET
    # ==========================================================

    @staticmethod
    def _add_bullet(
        document: Document,
        text: str,
    ) -> None:

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.left_indent = Inches(
            0.14
        )

        paragraph.paragraph_format.first_line_indent = Inches(
            -0.10
        )

        paragraph.paragraph_format.space_after = Pt(1)

        run = paragraph.add_run(
            "• "
        )

        run.font.name = "Arial"
        run.font.size = Pt(8.5)

        run = paragraph.add_run(
            text
        )

        run.font.name = "Arial"
        run.font.size = Pt(8.5)

    # ==========================================================
    # CAREER OBJECTIVE
    # ==========================================================

    @staticmethod
    def _create_objective(
        resume: Dict,
    ) -> str:

        summary = resume.get(
            "summary",
            "",
        ).strip()

        if summary:
            return summary

        return (
            "To secure a challenging position where I can "
            "apply my technical and analytical skills to "
            "solve business problems while contributing "
            "to organizational growth."
        )

    # ==========================================================
    # TWO-COLUMN WORD XML
    # ==========================================================

    @staticmethod
    def _enable_two_columns(
        document: Document,
    ) -> None:

        section = document.sections[0]

        sect_pr = section._sectPr

        columns = sect_pr.find(
            qn("w:cols")
        )

        if columns is None:

            columns = OxmlElement(
                "w:cols"
            )

            sect_pr.append(
                columns
            )

        columns.set(
            qn("w:num"),
            "2",
        )

        columns.set(
            qn("w:space"),
            "720",
        )

    @staticmethod
    def _add_column_break(
        document: Document,
    ) -> None:

        paragraph = document.add_paragraph()

        run = paragraph.add_run()

        break_element = OxmlElement(
            "w:br"
        )

        break_element.set(
            qn("w:type"),
            "column",
        )

        run._r.append(
            break_element
        )