from pathlib import Path
from typing import Union

import fitz
from docx import Document


class ResumeExtractor:
    """
    Extracts text from supported resume files.

    Supported formats:
    - PDF
    - DOCX
    """

    ALLOWED_EXTENSIONS = {
        ".pdf",
        ".docx",
    }

    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB

    @classmethod
    def extract(
        cls,
        file_path: Union[str, Path],
    ) -> str:
        """
        Extract text from a PDF or DOCX file.
        """

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Resume file not found: {path}"
            )

        extension = path.suffix.lower()

        if extension not in cls.ALLOWED_EXTENSIONS:
            raise ValueError(
                "Unsupported resume format. "
                "Only PDF and DOCX files are supported."
            )

        file_size = path.stat().st_size

        if file_size > cls.MAX_FILE_SIZE:
            raise ValueError(
                "Resume file is too large. "
                "Maximum allowed size is 5 MB."
            )

        if extension == ".pdf":
            text = cls._extract_pdf(path)

        elif extension == ".docx":
            text = cls._extract_docx(path)

        else:
            raise ValueError(
                "Unsupported resume format."
            )

        cleaned_text = cls.clean_text(text)

        if not cleaned_text:
            raise ValueError(
                "Could not extract readable text from the resume."
            )

        return cleaned_text

    @staticmethod
    def _extract_pdf(
        path: Path,
    ) -> str:
        """
        Extract text from all PDF pages.
        """

        pages = []

        try:
            document = fitz.open(path)

            try:
                for page in document:
                    page_text = page.get_text("text")

                    if page_text:
                        pages.append(page_text)

            finally:
                document.close()

        except Exception as exc:
            raise ValueError(
                f"Failed to read PDF resume: {exc}"
            ) from exc

        return "\n".join(pages)

    @staticmethod
    def _extract_docx(
        path: Path,
    ) -> str:
        """
        Extract paragraphs and table content from DOCX.
        """

        try:
            document = Document(path)

        except Exception as exc:
            raise ValueError(
                f"Failed to read DOCX resume: {exc}"
            ) from exc

        sections = []

        # Normal paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                sections.append(text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        cells.append(text)

                if cells:
                    sections.append(" | ".join(cells))

        return "\n".join(sections)

    @staticmethod
    def clean_text(
        text: str,
    ) -> str:
        """
        Clean extracted resume text while preserving
        useful line structure for section detection.
        """

        if not text:
            return ""

        # Normalize line endings.
        text = text.replace("\r\n", "\n")
        text = text.replace("\r", "\n")

        cleaned_lines = []

        for line in text.splitlines():

            # Replace non-breaking spaces.
            line = line.replace("\xa0", " ")

            # Collapse repeated spaces/tabs.
            line = " ".join(line.split())

            line = line.strip()

            if line:
                cleaned_lines.append(line)

        # Preserve line-based resume structure.
        return "\n".join(cleaned_lines).strip()