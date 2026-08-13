from pathlib import Path

from docx import Document


class DOCXParser:
    """
    Extracts text from DOCX resume files.

    Supports:
        - Normal paragraphs
        - Tables
    """

    @staticmethod
    def extract(file_path: str) -> str:
        """
        Extract text from paragraphs and tables in a DOCX file.

        Raises:
            FileNotFoundError:
                If the file does not exist.

            ValueError:
                If the file is invalid, unreadable, or contains no
                extractable text.
        """

        if not file_path:
            raise ValueError(
                "Resume file path is required."
            )

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Resume file not found: {file_path}"
            )

        if not path.is_file():
            raise ValueError(
                "The provided resume path is not a file."
            )

        if path.suffix.lower() != ".docx":
            raise ValueError(
                "DOCXParser only supports DOCX files."
            )

        if path.stat().st_size == 0:
            raise ValueError(
                "The DOCX resume file is empty."
            )

        try:
            document = Document(str(path))
        except Exception as exc:
            raise ValueError(
                "Unable to read the DOCX resume. "
                "The file may be corrupted or invalid."
            ) from exc

        sections = []

        # ---------------------------------------------------------
        # Normal paragraphs
        # ---------------------------------------------------------

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                sections.append(text)

        # ---------------------------------------------------------
        # Tables
        # ---------------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    text = cell.text.strip()

                    if text:
                        cells.append(text)

                if cells:
                    sections.append(
                        " | ".join(cells)
                    )

        extracted_text = "\n".join(
            sections
        ).strip()

        if not extracted_text:
            raise ValueError(
                "No readable text was found in the DOCX resume."
            )

        return extracted_text